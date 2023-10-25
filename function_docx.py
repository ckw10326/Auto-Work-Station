from docx import Document
import os
# 使用範例
#replace_text_in_docx('檔案路徑.docx', 'abc', '123')


def read_docx_file(file_path):
    try:
        doc = Document(file_path)
        # 讀取文件內容
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return content
    except Exception as e:
        print("讀取文件時發生錯誤:", str(e))
        return None

def replace_keyword_in_docx(file_path, keyword, replacement):
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if keyword in paragraph.text:
                paragraph.text = paragraph.text.replace(keyword, replacement)
        
        # 替換標題中的關鍵字
        for section in doc.sections:
            header = section.header
            if header is not None:
                for paragraph in header.paragraphs:
                    if keyword in paragraph.text:
                        paragraph.text = paragraph.text.replace(keyword, replacement)
            
            footer = section.footer
            if footer is not None:
                for paragraph in footer.paragraphs:
                    if keyword in paragraph.text:
                        paragraph.text = paragraph.text.replace(keyword, replacement)
        
        # 儲存替換後的文件
        new_file_path = file_path.replace('.docx', '_.docx')
        doc.save(new_file_path)
        os.remove(file_path)
        print("關鍵字替換完成。")
        return new_file_path
    except Exception as e:
        print("處理文件時發生錯誤:", str(e))

read_docx_file("/workspaces/Auto-Work-Station/00dest/00source/套印_TPC-TC(C0)-CD-23-3019.docx")